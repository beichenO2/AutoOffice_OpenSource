"""python-docx bridge: reads JSON payload from stdin, writes .docx to argv[1]."""
import base64, json, sys
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

THEMES = {
    "academic": {"font": "Times New Roman", "title_size": 18, "body_size": 12, "heading_color": "000000"},
    "business": {"font": "Calibri", "title_size": 20, "body_size": 11, "heading_color": "1F4E79"},
    "minimal":  {"font": "Arial",   "title_size": 16, "body_size": 11, "heading_color": "333333"},
}

def build(payload: dict, out_path: str) -> None:
    cfg = THEMES.get(payload.get("theme", "minimal"), THEMES["minimal"])
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = cfg["font"]
    style.font.size = Pt(cfg["body_size"])

    if payload.get("locale") == "zh-CN":
        style.font.name = "SimSun"

    title_para = doc.add_heading(payload["title"], level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(cfg["title_size"])

    if payload.get("subtitle"):
        sub = doc.add_paragraph(payload["subtitle"])
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_parts = []
    if payload.get("author"):
        meta_parts.append(payload["author"])
    if payload.get("date"):
        meta_parts.append(payload["date"])
    if meta_parts:
        mp = doc.add_paragraph(" | ".join(meta_parts))
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if payload.get("toc"):
        doc.add_paragraph("（目录占位 — Word 打开后按 Ctrl+A F9 更新域）")

    for sec in payload.get("sections", []):
        level = sec.get("level", 1)
        h = doc.add_heading(sec["heading"], level=min(level, 3))
        color = RGBColor.from_string(cfg["heading_color"])
        for run in h.runs:
            run.font.color.rgb = color

        if sec.get("body"):
            doc.add_paragraph(sec["body"])

        if sec.get("chartPngBase64"):
            doc.add_picture(BytesIO(base64.b64decode(sec["chartPngBase64"])), width=Inches(6.2))

        tbl = sec.get("table")
        if tbl and tbl.get("headers"):
            headers = tbl["headers"]
            rows = tbl.get("rows", [])
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for i, h_text in enumerate(headers):
                table.rows[0].cells[i].text = h_text
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    if ci < len(headers):
                        table.rows[ri + 1].cells[ci].text = str(cell)

    if payload.get("headerText"):
        for section in doc.sections:
            header = section.header
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.text = payload["headerText"]
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if payload.get("footerText"):
        for section in doc.sections:
            footer = section.footer
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.text = payload["footerText"]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)

if __name__ == "__main__":
    data = json.load(sys.stdin)
    build(data, sys.argv[1])
